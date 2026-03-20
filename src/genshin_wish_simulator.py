#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
原神抽卡模拟器 - 完整策略版
================================
功能：模拟6.4下半限定角色卡池（丝柯克、爱可菲）与武器卡池（苍耀、香韵奏者）
      - 完全基于官方规则，软保底采用社区公认概率曲线
      - 目标预设与达成预测（基于社区期望值）
      - 武器卡池每次出金后暂停，提供策略选项（含“保持定轨继续”）
      - 无定轨时命定值始终为0，且增加确认提示
      - 目标达成自动暂停（即时询问，避免重复提醒）
      - 重复角色转化规则（第2-7次与第8次+区分）
      - 策略日志记录，用于最终报告分析
      - 导出Excel（含行着色、汇总指标）、四张独立PNG图表
      - 速度模式选择，终端彩色输出
      - 所有用户可见文本通过 i18n 模块管理，支持多语言
      - 限定物品获取详情（Excel/Word）基于真实抽数
      - 修复四星物品星辉显示
      - 优化武器池定轨交互
      - 图表自适应：五星TOP10动态宽度，四星TOP10动态高度，饼图动态标签
      - Word报告字体统一（微软雅黑→宋体回退）
      - 角色和武器池真实抽数独立计算，汇总表合并平均真实抽数
"""

import random
import time
import datetime
import os
import math
from collections import Counter

import pandas as pd
import numpy as np
import plotly.graph_objects as go
from plotly.subplots import make_subplots
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import load_workbook
from openpyxl.styles import Font, PatternFill
from openpyxl.worksheet.table import Table, TableStyleInfo

from i18n import init_i18n, t

# 初始化翻译模块（默认为中文）
init_i18n("zh-CN")

# ==================== ANSI 颜色定义 ====================
COLOR_GOLD = "\033[93m"      # 金色
COLOR_PURPLE = "\033[95m"    # 紫色
COLOR_RESET = "\033[0m"      # 重置颜色
BOLD = "\033[1m"             # 加粗

def color_text(text, star):
    """根据星级返回带颜色和加粗的文本"""
    if star == 5:
        return f"{BOLD}{COLOR_GOLD}{text}{COLOR_RESET}"
    elif star == 4:
        return f"{BOLD}{COLOR_PURPLE}{text}{COLOR_RESET}"
    else:
        return text

# ==================== 1. 卡池与物品配置 ====================
VERSION = "6.4下半"
START_DATE = "2026-03-17"
END_DATE = "2026-04-07"

# 角色卡池（C1: 丝柯克，C2: 爱可菲）
CHAR_BANNERS = {
    "C1": {"name": t("banner.c1_name"), "weapon": "单手剑", "element": "冰", "type": "主C"},
    "C2": {"name": t("banner.c2_name"), "weapon": "长柄武器", "element": "冰", "type": "辅助"}
}
FOUR_STAR_CHARS_UP = ["塔利雅", "坎蒂丝", "夏洛蒂"]

# 武器卡池（W1: 苍耀，W2: 香韵奏者）
WEAPON_BANNERS = {
    "W1": {"name": t("banner.w1_name"), "character": "丝柯克"},
    "W2": {"name": t("banner.w2_name"), "character": "爱可菲"}
}
FOUR_STAR_WEAPONS_UP = ["笛剑", "钟剑", "西风长枪", "西风秘典", "弓藏"]

# 常驻五星角色（8个）
STANDARD_CHARS = ["梦见月瑞希", "迪希雅", "提纳里", "刻晴", "莫娜", "七七", "迪卢克", "琴"]

# 常驻五星武器（10个）
STANDARD_WEAPONS = [
    "天空之刃", "风鹰剑", "狼的末路", "天空之傲", "和璞鸢",
    "天空之脊", "四风原典", "天空之卷", "阿莫斯之弓", "天空之翼"
]

# 所有四星角色（46个，不含御三家）
ALL_FOUR_STAR_CHARS = [
    "塔利雅", "夏洛蒂", "坎蒂丝", "叶洛亚", "雅珂达", "爱诺",
    "伊法", "伊安珊", "蓝砚", "欧洛伦", "卡齐娜", "赛索斯",
    "嘉明", "夏沃蕾", "菲米尼", "琳妮特", "卡维", "米卡",
    "瑶瑶", "珐露珊", "莱依拉", "多莉", "柯莱", "久岐忍",
    "云堇", "绮良良", "鹿野院平藏", "九条裟罗", "五郎", "早柚",
    "托马", "烟绯", "罗莎莉亚", "辛焱", "砂糖", "迪奥娜",
    "重云", "诺艾尔", "班尼特", "菲谢尔", "凝光", "行秋",
    "北斗", "香菱", "雷泽", "芭芭拉"
]

# 所有四星武器（18把）
ALL_FOUR_STAR_WEAPONS = [
    "匣里龙吟", "祭礼剑", "笛剑", "西风剑",
    "雨裁", "祭礼大剑", "钟剑", "西风大剑",
    "西风长枪", "匣里灭辰",
    "昭心", "祭礼残章", "流浪乐章", "西风秘典",
    "弓藏", "祭礼弓", "绝弦", "西风猎弓"
]

# 三星武器（13种）
THREE_STAR_WEAPONS = [
    "飞天御剑", "黎明神剑", "冷刃",
    "以理服人", "沐浴龙血的剑", "铁影阔剑",
    "黑缨枪",
    "翡玉法球", "讨龙英杰谭", "魔导绪论",
    "弹弓", "神射手之誓", "鸦羽弓"
]

# ==================== 社区期望值数据 ====================
EXPECTATION = {
    'single_up_char': 93.36,      # 单UP角色期望抽数
    'max_pity_char': 180,          # 大保底上限
    'p75_char': 135,               # 75%成功率抽数
    'p90_char': 155,               # 90%成功率抽数
    'c6_char': 656,                # 满命角色期望抽数
    'c6_p75': 640,                 # 满命75%概率
    'c6_p90': 700,                 # 满命90%概率
    'single_up_weapon': 135,       # 单UP武器期望抽数
    'weapon_p75': 160,             # 武器75%概率
    'weapon_p90': 180,             # 武器90%概率
    'weapon_max': 240,             # 武器理论最大
    'dual_weapon_exp': 158.5,      # 双限定武器各一把期望
}

# ==================== 2. 抽卡状态管理类 ====================
class GachaState:
    def __init__(self):
        # 角色卡池状态
        self.char_pity = 0
        self.char_guarantee = False          # 五星大保底
        self.char_lost_streak = 0            # 连续歪次数
        self.char_four_star_guarantee = False # 四星UP保底（角色池）

        # 武器卡池状态
        self.weapon_pity = 0
        self.weapon_fate = 0
        self.weapon_chosen = None
        self.weapon_four_star_guarantee = False # 四星UP保底（武器池）

        # 共享四星保底计数器
        self.four_star_pity = 0

        self.total_draws = 0          # 总抽卡次数
        self.total_char_draws = 0     # 角色池累计抽数
        self.total_weapon_draws = 0   # 武器池累计抽数
        self.starglitter = 0
        # 角色计数（包括限定和常驻）
        self.char_count = {"丝柯克": 0, "爱可菲": 0}
        self.four_star_char_count = {}

        self.records_char = []       # 角色池记录
        self.records_weapon = []     # 武器池记录
        self.records_all = []         # 所有记录

        # 目标预设相关
        self.targets = {'C1': 0, 'C2': 0, 'W1': 0, 'W2': 0}  # 目标数量，0表示不限
        self.dual_weapon_mode = False        # 是否启用双限定武器策略模式
        self.strategy_log = []                # 策略日志

        # 记录每个限定物品上次获得时的各池累计抽数，用于计算真实抽数
        self.last_char_total = {"丝柯克": 0, "爱可菲": 0}
        self.last_weapon_total = {"苍耀": 0, "香韵奏者": 0}
        self.real_spins_limited = {"丝柯克": [], "爱可菲": [], "苍耀": [], "香韵奏者": []}

    def reset_weapon_fate(self):
        self.weapon_fate = 0
        self.weapon_chosen = None

# ==================== 3. 概率计算函数（社区公认模型）====================
def get_5star_prob(pity, banner_type):
    if banner_type == 'char':
        if pity <= 73:
            return 0.006
        elif pity == 74: return 0.066
        elif pity == 75: return 0.126
        elif pity == 76: return 0.186
        elif pity == 77: return 0.246
        elif pity == 78: return 0.306
        elif pity == 79: return 0.366
        elif pity == 80: return 0.426
        elif pity == 81: return 0.486
        elif pity == 82: return 0.546
        elif pity == 83: return 0.606
        elif pity == 84: return 0.666
        elif pity == 85: return 0.726
        elif pity == 86: return 0.786
        elif pity == 87: return 0.846
        elif pity == 88: return 0.906
        elif pity == 89: return 0.966
        else: return 1.0
    else:  # weapon
        if pity <= 62:
            return 0.007
        elif pity == 63: return 0.077
        elif pity == 64: return 0.147
        elif pity == 65: return 0.217
        elif pity == 66: return 0.287
        elif pity == 67: return 0.357
        elif pity == 68: return 0.427
        elif pity == 69: return 0.497
        elif pity == 70: return 0.567
        elif pity == 71: return 0.637
        elif pity == 72: return 0.707
        elif pity == 73: return 0.777
        elif pity == 74: return 0.847
        elif pity == 75: return 0.882
        elif pity == 76: return 0.917
        elif pity == 77: return 0.952
        elif pity == 78: return 0.987
        elif pity == 79: return 0.99
        else: return 1.0

def is_five_star(pity, banner_type):
    return random.random() < get_5star_prob(pity, banner_type)

# ==================== 4. 星辉计算（包含重复转化）====================
def starglitter_for_char(item_name, state, count):
    """计算五星角色获得的星辉"""
    if count <= 7:
        return 10
    else:
        return 25

def starglitter_for_four_star_char(item_name, state):
    """四星角色星辉，并更新计数"""
    if item_name in state.four_star_char_count:
        count = state.four_star_char_count[item_name] + 1
    else:
        count = 1
    state.four_star_char_count[item_name] = count
    return 2 if count <= 7 else 5

# ==================== 5. 单抽核心函数 ====================
def draw_one(state, banner_type, banner_code, effective_choice=None):
    draw_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    if banner_type == 'char':
        current_pity = state.char_pity + 1
    else:
        current_pity = state.weapon_pity + 1

    if banner_type == 'char':
        state.char_pity += 1
        pity = state.char_pity
        guarantee = state.char_guarantee
        lost_streak = state.char_lost_streak
        # 更新角色池累计抽数
        state.total_char_draws += 1
    else:
        state.weapon_pity += 1
        pity = state.weapon_pity
        # 更新武器池累计抽数
        state.total_weapon_draws += 1

    state.four_star_pity += 1
    state.total_draws += 1

    force_five = False
    if banner_type == 'char' and state.char_pity >= 90:
        force_five = True
    elif banner_type == 'weapon' and state.weapon_pity >= 80:
        force_five = True

    force_four = (state.four_star_pity >= 10)

    item_name = ""
    item_category = ""
    star = 0
    is_up = False
    is_capture = False
    guarantee_type = ""
    glitter = 0
    dust = 0
    is_overflow = False

    if force_five or is_five_star(pity, banner_type):
        if banner_type == 'char':
            # 角色池五星
            if guarantee:
                item_name = CHAR_BANNERS[banner_code]["name"]
                is_up = True
                guarantee_type = "guaranteed"
                state.char_guarantee = False
                state.char_lost_streak = 0
            else:
                capture_prob = 0.00018
                base_up_prob = [0.5, 0.55, 0.75][min(lost_streak, 2)]
                r = random.random()
                if r < capture_prob:
                    item_name = CHAR_BANNERS[banner_code]["name"]
                    is_up = True
                    is_capture = True
                    guarantee_type = "capturing"
                    state.char_lost_streak = 0
                elif r < capture_prob + base_up_prob:
                    item_name = CHAR_BANNERS[banner_code]["name"]
                    is_up = True
                    guarantee_type = "won_up"
                    state.char_lost_streak = 0
                else:
                    item_name = random.choice(STANDARD_CHARS)
                    is_up = False
                    guarantee_type = "lost"
                    state.char_guarantee = True
                    state.char_lost_streak += 1

            if item_name in state.char_count:
                old_count = state.char_count[item_name]
                new_count = old_count + 1
                state.char_count[item_name] = new_count
                glitter = starglitter_for_char(item_name, state, new_count)
                if new_count > 7:
                    is_overflow = True
            else:
                state.char_count[item_name] = 1
                glitter = 0

            item_category = t("category.character")
            star = 5
            state.char_pity = 0

            # 记录限定角色的真实抽数（基于角色池累计）
            if item_name in ["丝柯克", "爱可菲"]:
                real_spin = state.total_char_draws - state.last_char_total[item_name]
                state.real_spins_limited[item_name].append(real_spin)
                state.last_char_total[item_name] = state.total_char_draws

        else:  # weapon
            # 武器池五星
            if effective_choice is None:
                # 无定轨模式
                if random.random() < 0.75:
                    item_name = random.choice([WEAPON_BANNERS["W1"]["name"], WEAPON_BANNERS["W2"]["name"]])
                    guarantee_type = "won_up"
                    is_up = True
                else:
                    item_name = random.choice(STANDARD_WEAPONS)
                    guarantee_type = "lost"
                    is_up = False
                state.weapon_fate = 0
                state.weapon_chosen = None
            else:
                # 有定轨模式
                if state.weapon_fate >= 1:
                    item_name = WEAPON_BANNERS[state.weapon_chosen]["name"]
                    is_up = True
                    guarantee_type = "epitomized"
                    state.weapon_fate = 0
                else:
                    if random.random() < 0.75:
                        if random.random() < 0.5:
                            item_name = WEAPON_BANNERS[state.weapon_chosen]["name"]
                            guarantee_type = "won_up"
                        else:
                            other = [w for w in WEAPON_BANNERS if w != state.weapon_chosen][0]
                            item_name = WEAPON_BANNERS[other]["name"]
                            guarantee_type = "lost"
                            state.weapon_fate += 1
                    else:
                        item_name = random.choice(STANDARD_WEAPONS)
                        guarantee_type = "lost"
                        state.weapon_fate += 1
                    state.weapon_fate = min(state.weapon_fate, 1)
                    is_up = (item_name in [WEAPON_BANNERS[w]["name"] for w in WEAPON_BANNERS])

            item_category = t("category.weapon")
            star = 5
            glitter = 10
            state.weapon_pity = 0

            # 记录限定武器的真实抽数（基于武器池累计）
            if item_name in ["苍耀", "香韵奏者"]:
                real_spin = state.total_weapon_draws - state.last_weapon_total[item_name]
                state.real_spins_limited[item_name].append(real_spin)
                state.last_weapon_total[item_name] = state.total_weapon_draws

        state.four_star_pity = 0

    elif force_four or random.random() < (0.051 if banner_type == 'char' else 0.06):
        # 四星判定
        if banner_type == 'char':
            if state.char_four_star_guarantee:
                item_name = random.choice(FOUR_STAR_CHARS_UP)
                is_up = True
                state.char_four_star_guarantee = False
            else:
                r = random.random()
                if r < 0.5:
                    item_name = random.choice(FOUR_STAR_CHARS_UP)
                    is_up = True
                elif r < 0.75:
                    non_up_chars = [c for c in ALL_FOUR_STAR_CHARS if c not in FOUR_STAR_CHARS_UP]
                    item_name = random.choice(non_up_chars)
                    is_up = False
                    state.char_four_star_guarantee = True
                else:
                    item_name = random.choice(ALL_FOUR_STAR_WEAPONS)
                    is_up = False
                    state.char_four_star_guarantee = True

            # 判断物品类型并计算星辉
            if item_name in ALL_FOUR_STAR_CHARS or item_name in FOUR_STAR_CHARS_UP:
                item_category = t("category.character")
                glitter = starglitter_for_four_star_char(item_name, state)
                if state.four_star_char_count[item_name] > 7:
                    is_overflow = True
            else:
                item_category = t("category.weapon")
                glitter = 2

        else:  # weapon
            if state.weapon_four_star_guarantee:
                item_name = random.choice(FOUR_STAR_WEAPONS_UP)
                is_up = True
                state.weapon_four_star_guarantee = False
            else:
                r = random.random()
                if r < 0.75:
                    item_name = random.choice(FOUR_STAR_WEAPONS_UP)
                    is_up = True
                elif r < 0.875:
                    non_up_weapons = [w for w in ALL_FOUR_STAR_WEAPONS if w not in FOUR_STAR_WEAPONS_UP]
                    item_name = random.choice(non_up_weapons)
                    is_up = False
                    state.weapon_four_star_guarantee = True
                else:
                    item_name = random.choice(ALL_FOUR_STAR_CHARS)
                    is_up = False
                    state.weapon_four_star_guarantee = True

            # 判断物品类型并计算星辉
            if item_name in ALL_FOUR_STAR_CHARS or item_name in FOUR_STAR_CHARS_UP:
                item_category = t("category.character")
                glitter = starglitter_for_four_star_char(item_name, state)
            else:
                item_category = t("category.weapon")
                glitter = 2

        star = 4
        guarantee_type = "four_star"
        state.four_star_pity = 0

    else:
        # 三星武器
        item_name = random.choice(THREE_STAR_WEAPONS)
        item_category = t("category.weapon")
        star = 3
        guarantee_type = "three_star"
        dust = 15

    state.starglitter += glitter

    record = {
        "抽卡时间": draw_time,
        "卡池": banner_code,
        "获得物品": item_name,
        "类别": item_category,
        "星级": star,
        "总抽卡次数": state.total_draws,
        "小保底内次数": current_pity,
        "是否UP": is_up,
        "是否触发捕获明光": is_capture if star == 5 and banner_type=='char' else False,
        "保底类型": guarantee_type,
        "获得星辉": glitter,
        "获得星尘": dust,
        "备注": "溢出" if is_overflow else ""
    }

    if banner_type == 'char':
        state.records_char.append(record)
    else:
        record["定轨武器"] = state.weapon_chosen if state.weapon_chosen else ""
        record["命定值"] = state.weapon_fate
        state.records_weapon.append(record)

    state.records_all.append(record)
    return record

# ==================== 6. 目标预设与预测函数 ====================
def setup_targets():
    print(t("target.set_question"))
    choice = input().strip().upper()
    targets = {'C1': 0, 'C2': 0, 'W1': 0, 'W2': 0}
    dual_weapon = False
    if choice == 'Y':
        print(t("target.select_banners"))
        pools = input().strip().upper().replace('，', ',').split(',')
        pools = [p.strip() for p in pools]
        for p in pools:
            if p in ['C1', 'C2']:
                while True:
                    try:
                        val = int(input(t("target.input_char", banner=p, name=CHAR_BANNERS[p]['name'])))
                        if val >= 0:
                            targets[p] = val
                            break
                        else:
                            print(t("common.non_negative"))
                    except ValueError:
                        print(t("common.enter_number"))
            elif p == 'W':
                print(t("target.input_weapon"))
                while True:
                    try:
                        w1 = int(input(t("target.input_w1")))
                        if w1 >= 0:
                            targets['W1'] = w1
                            break
                        else:
                            print(t("common.non_negative"))
                    except ValueError:
                        print(t("common.enter_number"))
                while True:
                    try:
                        w2 = int(input(t("target.input_w2")))
                        if w2 >= 0:
                            targets['W2'] = w2
                            break
                        else:
                            print(t("common.non_negative"))
                    except ValueError:
                        print(t("common.enter_number"))
                if targets['W1'] > 0 and targets['W2'] > 0:
                    d = input(t("target.dual_question")).strip().upper()
                    dual_weapon = (d == 'Y')
        show_prediction(targets, dual_weapon)
        input(t("target.press_any"))
    return targets, dual_weapon

def show_prediction(targets, dual_weapon):
    """显示目标达成预计抽数"""
    print(t("target.prediction_title"))
    desc_parts = []
    if targets['C1'] > 0:
        desc_parts.append(f"丝柯克 {targets['C1']} 个")
    if targets['C2'] > 0:
        desc_parts.append(f"爱可菲 {targets['C2']} 个")
    weapon_parts = []
    if targets['W1'] > 0:
        weapon_parts.append(f"苍耀 {targets['W1']} 把")
    if targets['W2'] > 0:
        weapon_parts.append(f"香韵奏者 {targets['W2']} 把")
    if weapon_parts:
        desc_parts.append("，".join(weapon_parts))
    print(t("target.your_targets", targets=", ".join(desc_parts)))
    print(t("target.expectation_based"))

    total_expected = 0
    for char in ['C1', 'C2']:
        if targets[char] > 0:
            num = targets[char]
            if num == 7:
                print(t("target.char_c6"))
                print(t("target.char_c6_p75"))
                print(t("target.char_c6_p90"))
                total_expected += EXPECTATION['c6_char']
            else:
                exp = EXPECTATION['single_up_char'] * num
                print(f"• {CHAR_BANNERS[char]['name']} {num} 个：期望 {exp:.0f} 抽")
                total_expected += exp
    for w in ['W1', 'W2']:
        if targets[w] > 0:
            num = targets[w]
            if num == 1:
                print(t("target.weapon_single"))
                print(t("target.weapon_p75"))
                print(t("target.weapon_p90"))
                total_expected += EXPECTATION['single_up_weapon']
            else:
                exp = EXPECTATION['single_up_weapon'] * num
                print(f"• {WEAPON_BANNERS[w]['name']} {num} 把：期望 {exp:.0f} 抽")
                total_expected += exp
    if dual_weapon:
        print(t("target.dual_weapon_note"))
        print(t("target.dual_weapon_exp"))
        print(t("target.dual_weapon_strategy"))
    print(t("target.total_expected", total=total_expected))
    p90_total = total_expected * 1.1
    print(t("target.p90_suggestion", p90=p90_total))

# ==================== 7. 限定物品汇总函数 ====================
def print_limited_summary(state):
    skk = state.char_count.get("丝柯克", 0)
    akf = state.char_count.get("爱可菲", 0)
    w1 = sum(1 for r in state.records_weapon if r["获得物品"] == t("banner.w1_name"))
    w2 = sum(1 for r in state.records_weapon if r["获得物品"] == t("banner.w2_name"))
    msg = t("status.limited_summary",
            skk=skk, skk_cons=max(0, skk-1),
            akf=akf, akf_cons=max(0, akf-1),
            w1=w1, w2=w2)
    if skk > 7 or akf > 7:
        msg += t("status.overflow")
    print(msg)

# ==================== 7.5 预设目标进度打印 ====================
def print_target_status(state):
    has_target = any(v > 0 for v in state.targets.values())
    if not has_target:
        return
    print("\n【预设目标进度】")
    for k, target in state.targets.items():
        if target == 0:
            continue
        if k.startswith('C'):
            name = CHAR_BANNERS[k]['name']
            current = state.char_count.get(name, 0)
            status = "✅ 已达成" if current >= target else f"{current}/{target}"
            print(f"  {name}：{status}")
        else:  # W1, W2
            name = WEAPON_BANNERS[k]['name']
            current = sum(1 for r in state.records_weapon if r["获得物品"] == name)
            status = "✅ 已达成" if current >= target else f"{current}/{target}"
            print(f"  {name}：{status}")
    print()

# ==================== 8. 检查目标达成（用于抽卡过程中）====================
def check_target_achieved_during(state, banner_code, disabled_pools):
    for k, target in state.targets.items():
        if target == 0:
            continue
        if k.startswith('C'):
            if k != banner_code:
                continue
            char_name = CHAR_BANNERS[k]['name']
            current = state.char_count.get(char_name, 0)
            if current >= target:
                print(t("target_check.achieved", name=char_name, current=current, target=target))
                print(t("target_check.goal_status", name=char_name, current=current, target=target))
                choice = input(t("target_check.continue")).strip().upper()
                if choice == 'N':
                    disabled_pools.add(k)
                    print_target_status(state)
                    return True
        else:
            # 武器池的目标检查在 handle_weapon_five_star 中处理
            pass
    return False

# ==================== 9. 抽卡执行函数 ====================
def perform_draws(state, count, banner_type, banner_code, weapon_choice=None, no_delay=False, disabled_pools=None):
    if count <= 0:
        return False

    if banner_code in disabled_pools:
        print(t("command.disabled_pool", banner=banner_code))
        return False

    if banner_type == 'weapon':
        # 确定本次抽卡实际使用的定轨
        if weapon_choice is None:
            # 未指定定轨：沿用已有的定轨（如果有）
            effective_choice = state.weapon_chosen
            if effective_choice is None:
                # 既无指定也无已设定轨，需要用户确认
                print("提示：当前未设置定轨。建议使用定轨以提高目标武器获取概率。")
                confirm = input("是否继续无定轨抽卡？(Y/N)：").strip().upper()
                if confirm != 'Y':
                    print("已取消本次抽卡。")
                    return False
        else:
            effective_choice = weapon_choice
            # 如果指定了新的定轨且与当前不同，重置命定值
            if effective_choice != state.weapon_chosen:
                state.reset_weapon_fate()
                state.weapon_chosen = effective_choice

        remaining = count
        i = 0
        while remaining > 0:
            if remaining >= 10:
                i += 1
                results = []
                for _ in range(10):
                    results.append(draw_one(state, banner_type, banner_code, effective_choice))
                print(t("pull.ten_pull", i=i, banner=banner_code))
                for j, r in enumerate(results, 1):
                    star_str = t("pull.rarity", star=r['星级'])
                    name = r["获得物品"]
                    colored_name = color_text(name, r['星级'])
                    extra = t(f"pull.{r['保底类型']}") if r['保底类型'] in ['guaranteed','capturing','won_up','lost','epitomized'] else ""
                    glitter_info = t("pull.starglitter", count=r['获得星辉']) if r['获得星辉'] > 0 else ""
                    dust_info = t("pull.stardust", count=r['获得星尘']) if r['获得星尘'] > 0 else ""
                    overflow = t("pull.overflow") if r['备注'] == "溢出" else ""
                    print(f"  {j:2d}. {star_str} {colored_name}{extra}{glitter_info}{dust_info}{overflow}")
                if not no_delay:
                    time.sleep(1)
                remaining -= 10

                stop_command = False
                for r in results:
                    if r['星级'] == 5:
                        if handle_weapon_five_star(state, banner_code, effective_choice, remaining, disabled_pools):
                            stop_command = True
                            break
                if stop_command:
                    break
            else:
                r = draw_one(state, banner_type, banner_code, effective_choice)
                star_str = t("pull.rarity", star=r['星级'])
                name = r["获得物品"]
                colored_name = color_text(name, r['星级'])
                extra = t(f"pull.{r['保底类型']}") if r['保底类型'] in ['guaranteed','capturing','won_up','lost','epitomized'] else ""
                glitter_info = t("pull.starglitter", count=r['获得星辉']) if r['获得星辉'] > 0 else ""
                dust_info = t("pull.stardust", count=r['获得星尘']) if r['获得星尘'] > 0 else ""
                overflow = t("pull.overflow") if r['备注'] == "溢出" else ""
                print(f"{t('pull.pull_number', num=count - remaining + 1)}: {star_str} {colored_name}{extra}{glitter_info}{dust_info}{overflow}")
                remaining -= 1

                if r['星级'] == 5:
                    if handle_weapon_five_star(state, banner_code, effective_choice, remaining, disabled_pools):
                        break

                if (count - remaining) % 10 == 0 and remaining > 0 and not no_delay:
                    print_status(state)
                    time.sleep(1)
        if remaining == 0:
            print_status(state)
        return False
    else:
        # 角色池
        if count % 10 == 0:
            num_tens = count // 10
            for i in range(num_tens):
                results = []
                for _ in range(10):
                    results.append(draw_one(state, banner_type, banner_code, None))
                print(t("pull.ten_pull", i=i+1, banner=banner_code))
                for j, r in enumerate(results, 1):
                    star_str = t("pull.rarity", star=r['星级'])
                    name = r["获得物品"]
                    colored_name = color_text(name, r['星级'])
                    extra = t(f"pull.{r['保底类型']}") if r['保底类型'] in ['guaranteed','capturing','won_up','lost','epitomized'] else ""
                    glitter_info = t("pull.starglitter", count=r['获得星辉']) if r['获得星辉'] > 0 else ""
                    dust_info = t("pull.stardust", count=r['获得星尘']) if r['获得星尘'] > 0 else ""
                    overflow = t("pull.overflow") if r['备注'] == "溢出" else ""
                    print(f"  {j:2d}. {star_str} {colored_name}{extra}{glitter_info}{dust_info}{overflow}")
                if not no_delay:
                    time.sleep(1)

                for r in results:
                    if r['星级'] == 5:
                        if check_target_achieved_during(state, banner_code, disabled_pools):
                            print_status(state)
                            return True
                if (i + 1) % 5 == 0 or i == num_tens - 1:
                    print_status(state)
        else:
            print(t("pull.single_start", count=count, banner=banner_code))
            for i in range(1, count + 1):
                r = draw_one(state, banner_type, banner_code, None)
                star_str = t("pull.rarity", star=r['星级'])
                name = r["获得物品"]
                colored_name = color_text(name, r['星级'])
                extra = t(f"pull.{r['保底类型']}") if r['保底类型'] in ['guaranteed','capturing','won_up','lost','epitomized'] else ""
                glitter_info = t("pull.starglitter", count=r['获得星辉']) if r['获得星辉'] > 0 else ""
                dust_info = t("pull.stardust", count=r['获得星尘']) if r['获得星尘'] > 0 else ""
                overflow = t("pull.overflow") if r['备注'] == "溢出" else ""
                print(f"{t('pull.pull_number', num=i)}: {star_str} {colored_name}{extra}{glitter_info}{dust_info}{overflow}")

                if r['星级'] == 5:
                    if check_target_achieved_during(state, banner_code, disabled_pools):
                        print_status(state)
                        return True

                if i % 10 == 0:
                    print_status(state)
                    if i < count and not no_delay:
                        time.sleep(1)

            if count % 10 != 0:
                print_status(state)

        return False

def handle_weapon_five_star(state, banner_code, weapon_choice, remaining, disabled_pools):
    print_status(state)
    chosen_name = WEAPON_BANNERS[weapon_choice]["name"] if weapon_choice else t("weapon.no_path")
    w1 = sum(1 for r in state.records_weapon if r["获得物品"] == t("banner.w1_name"))
    w2 = sum(1 for r in state.records_weapon if r["获得物品"] == t("banner.w2_name"))
    fate = state.weapon_fate

    print(t("weapon.five_star_title", chosen=chosen_name, fate=fate))
    print(t("weapon.obtained", w1=w1, w2=w2))
    print(t("weapon.remaining", remaining=remaining))

    # 检查武器目标是否达成
    target_w1 = state.targets.get('W1', 0)
    target_w2 = state.targets.get('W2', 0)
    w1_achieved = target_w1 > 0 and w1 >= target_w1
    w2_achieved = target_w2 > 0 and w2 >= target_w2

    if (target_w1 > 0 and w1_achieved) or (target_w2 > 0 and w2_achieved):
        achieved_weapons = []
        if w1_achieved:
            achieved_weapons.append(t("banner.w1_name"))
        if w2_achieved:
            achieved_weapons.append(t("banner.w2_name"))
        print(f"\n🎯 目标达成：{'、'.join(achieved_weapons)}！")
        cont = input("是否继续抽取武器池？(Y/N)：").strip().upper()
        if cont != 'Y':
            disabled_pools.add('W')
            return True

    dual_active = state.dual_weapon_mode and (target_w1 > 0 and target_w2 > 0) and (not w1_achieved or not w2_achieved)

    options = []
    options.append(("1", t("weapon.option1")))
    options.append(("2", t("weapon.option2")))
    if weapon_choice:
        options.append(("3", t("weapon.option3")))
        options.append(("4", t("weapon.option4")))
    if dual_active and fate == 1 and weapon_choice:
        options.append(("5", t("weapon.option5")))

    print(t("weapon.choose_action"))
    for key, desc in options:
        print(f"{key}. {desc}")

    choice = input().strip()
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    if choice == '1':
        state.strategy_log.append(("continue", timestamp, None, None))
        state.reset_weapon_fate()
        weapon_choice = None
        return False
    elif choice == '2':
        state.strategy_log.append(("stop", timestamp, None, None))
        return True
    elif choice == '3' and weapon_choice:
        print(t("weapon.enter_new"))
        new_wc = input().strip().upper()
        if new_wc in ['W1', 'W2']:
            state.strategy_log.append(("change", timestamp, weapon_choice, new_wc))
            weapon_choice = new_wc
            state.reset_weapon_fate()
            state.weapon_chosen = new_wc
        else:
            print(t("weapon.invalid"))
        return False
    elif choice == '4' and weapon_choice:
        state.strategy_log.append(("keep", timestamp, None, None))
        return False
    elif choice == '5' and dual_active and fate == 1 and weapon_choice:
        state.strategy_log.append(("strategy", timestamp, weapon_choice, None))
        state.reset_weapon_fate()
        state.weapon_chosen = weapon_choice
        return False
    else:
        print(t("weapon.default"))
        state.reset_weapon_fate()
        weapon_choice = None
        return False

def print_status(state):
    total_dust = sum(r['获得星尘'] for r in state.records_all)
    c1_cnt = sum(1 for r in state.records_char if r['卡池'] == 'C1')
    c2_cnt = sum(1 for r in state.records_char if r['卡池'] == 'C2')
    w_cnt = len(state.records_weapon)
    print(t("status.title", total=state.total_draws, c1=c1_cnt, c2=c2_cnt, w=w_cnt,
             primogems=state.total_draws*160, glitter=state.starglitter, dust=total_dust))

# ==================== 10. 指标计算函数 ====================
def calculate_metrics(state):
    char_up_spins = []
    weapon_up_spins = []
    char_all_spins = []
    weapon_all_spins = []

    for r in state.records_char:
        if r['星级'] == 5:
            char_all_spins.append(r['小保底内次数'])
            if r['是否UP']:
                char_up_spins.append(r['小保底内次数'])

    for r in state.records_weapon:
        if r['星级'] == 5:
            weapon_all_spins.append(r['小保底内次数'])
            if r['是否UP']:
                weapon_up_spins.append(r['小保底内次数'])

    avg_char = sum(char_all_spins)/len(char_all_spins) if char_all_spins else 0
    avg_weapon = sum(weapon_all_spins)/len(weapon_all_spins) if weapon_all_spins else 0

    consecutive = 0
    max_consecutive = 0
    for r in sorted(state.records_char, key=lambda x: x['总抽卡次数']):
        if r['星级'] == 5:
            if r['是否UP']:
                consecutive += 1
                max_consecutive = max(max_consecutive, consecutive)
            else:
                consecutive = 0

    capture_total = sum(1 for r in state.records_char if r.get('是否触发捕获明光'))
    early_capture = capture_total

    total_char_5 = len([r for r in state.records_char if r['星级'] == 5])
    up_char = len([r for r in state.records_char if r['星级'] == 5 and r['是否UP']])
    no_loss_rate = up_char / total_char_5 if total_char_5 > 0 else 0

    fate_values = [r['命定值'] for r in state.records_weapon if r['星级'] == 5 and r['保底类型'] == 'epitomized']
    avg_fate = sum(fate_values)/len(fate_values) if fate_values else 0

    early_gold = 0
    for r in state.records_all:
        if r['星级'] == 5 and r['小保底内次数'] is not None:
            if (r['卡池'] in ['C1','C2'] and r['小保底内次数'] <= 40) or (r['卡池'] == 'W' and r['小保底内次数'] <= 35):
                early_gold += 1

    all_5 = [r for r in state.records_all if r['星级'] == 5]
    all_5_sorted = sorted(all_5, key=lambda x: x['总抽卡次数'])
    consecutive_gold = 0
    for i in range(len(all_5_sorted)-1):
        if all_5_sorted[i+1]['总抽卡次数'] - all_5_sorted[i]['总抽卡次数'] <= 20:
            consecutive_gold += 1

    multi_gold_10 = 0
    all_records_sorted = sorted(state.records_all, key=lambda x: x['总抽卡次数'])
    group = []
    for r in all_records_sorted:
        group.append(r)
        if len(group) == 10:
            golds = sum(1 for g in group if g['星级'] == 5)
            if golds >= 2:
                multi_gold_10 += 1
            group = []

    total_char_draws = len(state.records_char)
    real_up_char = total_char_draws / up_char if up_char > 0 else 0

    total_weapon_draws = len(state.records_weapon)
    weapon_up_total = len([r for r in state.records_weapon if r['星级'] == 5 and r['是否UP']])
    real_up_weapon = total_weapon_draws / weapon_up_total if weapon_up_total > 0 else 0

    e_score = 0
    if avg_char < 50:
        e_score += 30
    elif avg_char < 60:
        e_score += 20
    elif avg_char < 70:
        e_score += 10
    elif avg_char >= 75:
        e_score -= 10
    if avg_weapon < 55:
        e_score += 20
    elif avg_weapon < 65:
        e_score += 10
    elif avg_weapon >= 75:
        e_score -= 10

    a_score = 0
    if no_loss_rate >= 0.7:
        a_score += 20
    elif no_loss_rate >= 0.6:
        a_score += 15
    elif no_loss_rate >= 0.5:
        a_score += 10
    elif no_loss_rate < 0.3:
        a_score -= 20
    elif no_loss_rate < 0.4:
        a_score -= 10
    if avg_fate <= 0.5:
        a_score += 20
    elif avg_fate <= 1.0:
        a_score += 15
    elif avg_fate <= 1.5:
        a_score += 10
    elif avg_fate >= 2.0:
        a_score -= 10

    l_score = (capture_total * 3) + (multi_gold_10 * 5) + (early_gold * 2) + (consecutive_gold * 2)

    total_score = e_score * 0.4 + a_score * 0.4 + l_score * 0.2

    eggs = []
    soft_pity_count = sum(1 for s in char_all_spins if 74 <= s <= 78)
    if len(char_all_spins) > 0 and soft_pity_count / len(char_all_spins) > 0.8:
        eggs.append(t("eggs.soft_pity"))
    overflow_count = sum(1 for k, v in state.char_count.items() if k in ["丝柯克","爱可菲"] and v > 7)
    if overflow_count > 2:
        eggs.append(t("eggs.inventory_victim"))
    if avg_fate < 1.0 and (early_gold > 0 or multi_gold_10 > 0):
        eggs.append(t("eggs.weapon_master"))
    four_star_overflow = sum(1 for v in state.four_star_char_count.values() if v >= 8)
    if four_star_overflow >= 5:
        eggs.append(t("eggs.four_star_collector"))

    target_achieved = True
    for k, v in state.targets.items():
        if k.startswith('C'):
            char_name = CHAR_BANNERS[k]['name']
            if state.char_count.get(char_name, 0) < v:
                target_achieved = False
        else:
            weapon_name = WEAPON_BANNERS[k]['name']
            count = sum(1 for r in state.records_weapon if r['获得物品'] == weapon_name)
            if count < v:
                target_achieved = False

    five_star_dup_2_7 = 0
    five_star_dup_8plus = 0
    for name, cnt in state.char_count.items():
        if name in ["丝柯克", "爱可菲"] or name in STANDARD_CHARS:
            if cnt >= 2:
                if cnt <= 7:
                    five_star_dup_2_7 += cnt - 1
                else:
                    five_star_dup_2_7 += 6
                    five_star_dup_8plus += cnt - 7

    # 使用真实抽数计算平均（各物品独立）
    avg_skk = sum(state.real_spins_limited["丝柯克"]) / len(state.real_spins_limited["丝柯克"]) if state.real_spins_limited["丝柯克"] else 0
    avg_akf = sum(state.real_spins_limited["爱可菲"]) / len(state.real_spins_limited["爱可菲"]) if state.real_spins_limited["爱可菲"] else 0
    avg_w1 = sum(state.real_spins_limited["苍耀"]) / len(state.real_spins_limited["苍耀"]) if state.real_spins_limited["苍耀"] else 0
    avg_w2 = sum(state.real_spins_limited["香韵奏者"]) / len(state.real_spins_limited["香韵奏者"]) if state.real_spins_limited["香韵奏者"] else 0

    # 合并平均真实抽数（用于汇总表）
    all_char_spins = state.real_spins_limited["丝柯克"] + state.real_spins_limited["爱可菲"]
    all_weapon_spins = state.real_spins_limited["苍耀"] + state.real_spins_limited["香韵奏者"]
    avg_all_char = sum(all_char_spins) / len(all_char_spins) if all_char_spins else 0
    avg_all_weapon = sum(all_weapon_spins) / len(all_weapon_spins) if all_weapon_spins else 0

    return {
        'avg_up_char': round(avg_char, 1),
        'avg_up_weapon': round(avg_weapon, 1),
        'max_consecutive_up': max_consecutive,
        'capture_total': capture_total,
        'early_capture': early_capture,
        'no_loss_rate': no_loss_rate,
        'early_gold': early_gold,
        'consecutive_gold': consecutive_gold,
        'multi_gold_10': multi_gold_10,
        'real_up_char': round(real_up_char, 1),
        'real_up_weapon': round(real_up_weapon, 1),
        'total_score': round(total_score, 1),
        'e_score': e_score,
        'a_score': a_score,
        'l_score': l_score,
        'eggs': eggs,
        'target_achieved': target_achieved,
        'five_star_dup_2_7': five_star_dup_2_7,
        'five_star_dup_8plus': five_star_dup_8plus,
        'avg_skk': round(avg_skk, 1),
        'avg_akf': round(avg_akf, 1),
        'avg_w1': round(avg_w1, 1),
        'avg_w2': round(avg_w2, 1),
        'avg_all_char': round(avg_all_char, 1),
        'avg_all_weapon': round(avg_all_weapon, 1),
        'real_spins_limited': state.real_spins_limited,
    }

def get_luck_level(score, total_5star):
    if total_5star == 0:
        return None, t("luck.insufficient_data_title"), t("luck.insufficient_data_desc")
    if score >= 80:
        return 9, t("luck.level9_title"), t("luck.level9_desc")
    elif score >= 60:
        return 8, t("luck.level8_title"), t("luck.level8_desc")
    elif score >= 40:
        return 7, t("luck.level7_title"), t("luck.level7_desc")
    elif score >= 20:
        return 6, t("luck.level6_title"), t("luck.level6_desc")
    elif score >= 0:
        return 5, t("luck.level5_title"), t("luck.level5_desc")
    elif score >= -20:
        return 4, t("luck.level4_title"), t("luck.level4_desc")
    elif score >= -40:
        return 3, t("luck.level3_title"), t("luck.level3_desc")
    elif score >= -60:
        return 2, t("luck.level2_title"), t("luck.level2_desc")
    else:
        return 1, t("luck.level1_title"), t("luck.level1_desc")

# ==================== 11. 保存Excel ====================
def save_to_excel(state, output_dir, expected_total=0):
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = os.path.join(output_dir, f"抽卡记录_{timestamp}.xlsx")
    os.makedirs(output_dir, exist_ok=True)

    metrics = calculate_metrics(state)

    with pd.ExcelWriter(filename, engine='openpyxl') as writer:
        if state.records_char:
            df_char = pd.DataFrame(state.records_char)
            cols = ["抽卡时间", "卡池", "获得物品", "类别", "星级", "总抽卡次数", "小保底内次数",
                    "是否UP", "是否触发捕获明光", "保底类型", "获得星辉", "获得星尘", "备注"]
            df_char = df_char[[c for c in cols if c in df_char.columns]]
            df_char.to_excel(writer, sheet_name='角色记录', index=False)

        if state.records_weapon:
            df_weapon = pd.DataFrame(state.records_weapon)
            cols = ["抽卡时间", "卡池", "获得物品", "类别", "星级", "总抽卡次数", "小保底内次数",
                    "是否UP", "定轨武器", "命定值", "保底类型", "获得星辉", "获得星尘", "备注"]
            df_weapon = df_weapon[[c for c in cols if c in df_weapon.columns]]
            df_weapon.to_excel(writer, sheet_name='武器记录', index=False)

        # 新增：限定物品详情表（基于真实抽数）
        limited_rows = []
        for item, spins in metrics['real_spins_limited'].items():
            for idx, spin in enumerate(spins, 1):
                limited_rows.append({"物品": item, "第几次": idx, "真实抽数": spin})
        if limited_rows:
            df_limited = pd.DataFrame(limited_rows)
            df_limited.to_excel(writer, sheet_name='限定物品详情', index=False)

        expected = expected_total if expected_total else 0
        actual = state.total_draws
        deviation = actual - expected
        deviation_percent = (deviation / expected * 100) if expected > 0 else 0

        strategy_score = "优" if len(state.strategy_log) > 0 and metrics['multi_gold_10'] > 0 else "良"

        summary_data = {
            "统计项": [
                "总抽卡次数", "总获得星辉", "总获得星尘",
                "丝柯克数量", "爱可菲数量", "苍耀数量", "香韵奏者数量",
                "四星角色总数", "四星武器总数",
                "四星UP角色总数", "四星UP武器总数",
                "角色池平均UP抽数", "武器池平均UP抽数",
                "最多连续UP次数", "捕获明光次数", "提前触发捕获明光",
                "小保底不歪概率", "提前金次数", "连金次数", "十连多金次数",
                "限定角色平均真实抽数", "限定武器平均真实抽数",
                "综合得分", "运气等级", "彩蛋标识",
                "目标预设", "目标详情", "预计期望抽数", "实际消耗抽数", "偏差百分比",
                "策略执行次数", "目标达成情况", "策略效率评分",
                "五星角色转化(2-7次)", "五星角色转化(8次+)",
            ],
            "数值": [
                state.total_draws,
                state.starglitter,
                sum(r["获得星尘"] for r in state.records_all),
                state.char_count.get("丝柯克", 0),
                state.char_count.get("爱可菲", 0),
                sum(1 for r in state.records_weapon if r["获得物品"] == t("banner.w1_name")),
                sum(1 for r in state.records_weapon if r["获得物品"] == t("banner.w2_name")),
                sum(1 for r in state.records_all if r["星级"] == 4 and r["类别"] == "角色"),
                sum(1 for r in state.records_all if r["星级"] == 4 and r["类别"] == "武器"),
                sum(1 for r in state.records_all if r["星级"] == 4 and r["是否UP"] and r["类别"] == "角色"),
                sum(1 for r in state.records_all if r["星级"] == 4 and r["是否UP"] and r["类别"] == "武器"),
                metrics['avg_up_char'],
                metrics['avg_up_weapon'],
                metrics['max_consecutive_up'],
                metrics['capture_total'],
                metrics['early_capture'],
                f"{metrics['no_loss_rate']:.1%}",
                metrics['early_gold'],
                metrics['consecutive_gold'],
                metrics['multi_gold_10'],
                f"{metrics['avg_all_char']:.1f} 抽",
                f"{metrics['avg_all_weapon']:.1f} 抽",
                metrics['total_score'],
                get_luck_level(metrics['total_score'], sum(1 for r in state.records_all if r["星级"] == 5))[1],
                ", ".join(metrics['eggs']),
                "是" if any(state.targets.values()) else "否",
                str(state.targets),
                f"{expected:.0f}",
                actual,
                f"{deviation_percent:+.1f}%",
                len(state.strategy_log),
                "达成" if metrics['target_achieved'] else "未达成",
                strategy_score,
                metrics['five_star_dup_2_7'],
                metrics['five_star_dup_8plus'],
            ]
        }
        df_summary = pd.DataFrame(summary_data)
        df_summary.to_excel(writer, sheet_name='汇总', index=False)

        if state.strategy_log:
            log_entries = []
            for log in state.strategy_log:
                log_type, ts, arg1, arg2 = log
                if log_type == "continue":
                    desc = t("weapon.log_continue")
                elif log_type == "stop":
                    desc = t("weapon.log_stop")
                elif log_type == "change":
                    desc = t("weapon.log_change", old=arg1, new=arg2)
                elif log_type == "strategy":
                    desc = t("weapon.log_strategy", weapon=arg1)
                elif log_type == "keep":
                    desc = "保持定轨继续"
                else:
                    desc = str(log)
                log_entries.append({"决策时间": ts, "决策内容": desc})
            df_strategy = pd.DataFrame(log_entries)
            df_strategy.to_excel(writer, sheet_name='策略日志', index=False)

    wb = load_workbook(filename)
    msyh_font = Font(name='微软雅黑', size=11)
    gold_font = Font(name='微软雅黑', size=11, color='FFD700', bold=True)
    purple_font = Font(name='微软雅黑', size=11, color='800080', bold=True)

    for sheet_name in wb.sheetnames:
        if sheet_name in ['角色记录', '武器记录']:
            ws = wb[sheet_name]
            for row in ws.iter_rows():
                for cell in row:
                    cell.font = msyh_font
            star_col = None
            for col in range(1, ws.max_column + 1):
                if ws.cell(row=1, column=col).value == "星级":
                    star_col = col
                    break
            if star_col:
                for row in range(2, ws.max_row + 1):
                    star_val = ws.cell(row=row, column=star_col).value
                    if star_val == 5:
                        for col in range(1, ws.max_column + 1):
                            ws.cell(row=row, column=col).font = gold_font
                    elif star_val == 4:
                        for col in range(1, ws.max_column + 1):
                            ws.cell(row=row, column=col).font = purple_font
        else:
            ws = wb[sheet_name]
            for row in ws.iter_rows():
                for cell in row:
                    cell.font = msyh_font

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        table_ref = f"A1:{chr(64 + ws.max_column)}{ws.max_row}"
        table = Table(displayName=sheet_name.replace(" ", "_") + "表", ref=table_ref)
        style = TableStyleInfo(name="TableStyleLight1", showFirstColumn=False,
                               showLastColumn=False, showRowStripes=False, showColumnStripes=False)
        table.tableStyleInfo = style
        ws.add_table(table)

    wb.save(filename)
    print(t("export.excel_saved", filename=filename))
    return filename

# ==================== 12. 生成PNG图片 ====================
def generate_plots(state, excel_file):
    """生成四张独立的图表：角色五星饼图、武器五星饼图、四星TOP10、五星TOP10"""
    df_char = pd.DataFrame(state.records_char) if state.records_char else pd.DataFrame()
    df_weapon = pd.DataFrame(state.records_weapon) if state.records_weapon else pd.DataFrame()

    base = excel_file.replace(".xlsx", "")
    chart_files = []

    try:
        # 1. 角色池五星类型分布饼图（动态标签）
        if not df_char.empty:
            char_5 = df_char[df_char["星级"] == 5]
            if not char_5.empty:
                skk = sum(char_5["获得物品"] == "丝柯克")
                akf = sum(char_5["获得物品"] == "爱可菲")
                std = sum(char_5["获得物品"].isin(STANDARD_CHARS))
                
                # 构建动态标签和值
                labels = []
                values = []
                if skk > 0:
                    labels.append("丝柯克")
                    values.append(skk)
                if akf > 0:
                    labels.append("爱可菲")
                    values.append(akf)
                if std > 0:
                    labels.append("常驻角色")
                    values.append(std)
                
                if labels:
                    fig1 = go.Figure(data=[go.Pie(labels=labels, values=values)])
                    fig1.update_layout(title_text="角色卡池五星类型分布", font=dict(family="Microsoft YaHei", size=12))
                    fig1.update_traces(textposition='inside', textinfo='percent+label')
                    p1 = base + "_角色五星饼图.png"
                    fig1.write_image(p1)
                    chart_files.append(p1)
                    print(t("export.chart_saved", filename=p1))

        # 2. 武器池五星类型分布饼图（动态标签）
        if not df_weapon.empty:
            wp_5 = df_weapon[df_weapon["星级"] == 5]
            if not wp_5.empty:
                w1 = sum(wp_5["获得物品"] == t("banner.w1_name"))
                w2 = sum(wp_5["获得物品"] == t("banner.w2_name"))
                std_w = sum(wp_5["获得物品"].isin(STANDARD_WEAPONS))
                
                labels = []
                values = []
                if w1 > 0:
                    labels.append(t("banner.w1_name"))
                    values.append(w1)
                if w2 > 0:
                    labels.append(t("banner.w2_name"))
                    values.append(w2)
                if std_w > 0:
                    labels.append("常驻武器")
                    values.append(std_w)
                
                if labels:
                    fig2 = go.Figure(data=[go.Pie(labels=labels, values=values)])
                    fig2.update_layout(title_text="武器卡池五星类型分布", font=dict(family="Microsoft YaHei", size=12))
                    fig2.update_traces(textposition='inside', textinfo='percent+label')
                    p2 = base + "_武器五星饼图.png"
                    fig2.write_image(p2)
                    chart_files.append(p2)
                    print(t("export.chart_saved", filename=p2))

        # 3. 四星物品数量TOP10横向条形图（动态高度）
        four_star_items = Counter()
        for r in state.records_all:
            if r["星级"] == 4:
                four_star_items[r["获得物品"]] += 1
        top4 = four_star_items.most_common(10)
        if top4:
            names, counts = zip(*top4)
            num_items = len(names)
            # 动态高度：每个条目约40px，最小高度300px
            height = max(300, num_items * 40 + 80)  # 加80用于标题和边距
            fig3 = go.Figure(data=[go.Bar(x=counts, y=names, orientation='h')])
            fig3.update_layout(
                title_text="四星物品获取数量TOP10",
                font=dict(family="Microsoft YaHei", size=12),
                xaxis_title="数量",
                yaxis_title="物品",
                width=600,
                height=height,
                bargap=0.2,
                margin=dict(l=120, r=50, t=80, b=50)
            )
            fig3.update_yaxes(categoryorder='total descending')
            fig3.update_traces(texttemplate='%{x}', textposition='outside')
            p3 = base + "_四星TOP10.png"
            fig3.write_image(p3)
            chart_files.append(p3)
            print(t("export.chart_saved", filename=p3))

        # 4. 五星物品数量TOP10条形图（自动调整尺寸）
        five_star_items = Counter()
        for r in state.records_all:
            if r["星级"] == 5:
                five_star_items[r["获得物品"]] += 1
        top5 = five_star_items.most_common(10)
        if top5:
            names, counts = zip(*top5)
            num_items = len(names)
            width = max(400, num_items * 100)
            height = 400
            fig4 = go.Figure(data=[go.Bar(x=names, y=counts)])
            fig4.update_layout(
                title_text="五星物品获取数量TOP10",
                font=dict(family="Microsoft YaHei", size=12),
                xaxis_title="物品",
                yaxis_title="数量",
                width=width,
                height=height,
                bargap=0.3,
                margin=dict(l=50, r=50, t=80, b=50)
            )
            fig4.update_traces(
                texttemplate='%{y}',
                textposition='outside',
                textfont=dict(size=10)
            )
            p4 = base + "_五星TOP10.png"
            fig4.write_image(p4)
            chart_files.append(p4)
            print(t("export.chart_saved", filename=p4))

    except Exception as e:
        print(f"❌ 生成图表时发生错误: {e}")
        return None

    return chart_files

# ==================== 13. 生成Word报告 ====================
def set_run_font(run, font_name='微软雅黑', size=11, bold=False):
    """设置run的字体，优先使用微软雅黑，若系统不支持则尝试宋体"""
    try:
        run.font.name = font_name
    except:
        run.font.name = '宋体'
    run.font.size = Pt(size)
    run.bold = bold

def generate_word_report(state, excel_file, chart_files, expected_total=0):
    doc = Document()
    
    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    try:
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    except:
        pass

    title = doc.add_heading('原神抽卡分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_run_font(run, '微软雅黑', 16, bold=True)

    p = doc.add_paragraph(f"版本：{VERSION}")
    set_run_font(p.runs[0], '微软雅黑', 11)
    p = doc.add_paragraph(f"卡池时间：{START_DATE} 至 {END_DATE}")
    set_run_font(p.runs[0], '微软雅黑', 11)
    p = doc.add_paragraph(f"报告生成时间：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    set_run_font(p.runs[0], '微软雅黑', 11)

    metrics = calculate_metrics(state)
    total_5star = sum(1 for r in state.records_all if r["星级"] == 5)
    level, title_name, desc = get_luck_level(metrics['total_score'], total_5star)

    heading = doc.add_heading('1. 运气评定', level=1)
    for run in heading.runs:
        set_run_font(run, '微软雅黑', 14, bold=True)
    if level is None:
        p = doc.add_paragraph(title_name)
        set_run_font(p.runs[0], '微软雅黑', 11)
        p = doc.add_paragraph(desc)
        set_run_font(p.runs[0], '微软雅黑', 11)
    else:
        p = doc.add_paragraph(f"综合得分：{metrics['total_score']} 分")
        set_run_font(p.runs[0], '微软雅黑', 11)
        p = doc.add_paragraph(f"运气等级：Lv.{level} {title_name}")
        set_run_font(p.runs[0], '微软雅黑', 11)
        if metrics['eggs']:
            p = doc.add_paragraph(f"彩蛋标识：{', '.join(metrics['eggs'])}")
            set_run_font(p.runs[0], '微软雅黑', 11)
        p = doc.add_paragraph(desc)
        set_run_font(p.runs[0], '微软雅黑', 11)

    heading = doc.add_heading('2. 抽卡统计', level=1)
    for run in heading.runs:
        set_run_font(run, '微软雅黑', 14, bold=True)
    c1_count = sum(1 for r in state.records_all if r["卡池"] == "C1")
    c2_count = sum(1 for r in state.records_all if r["卡池"] == "C2")
    w_count = len(state.records_weapon)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, '微软雅黑', 11)
    table.cell(0, 0).text = "卡池"
    table.cell(0, 1).text = "抽卡次数"
    table.cell(1, 0).text = "C1（丝柯克）"
    table.cell(1, 1).text = str(c1_count)
    table.cell(2, 0).text = "C2（爱可菲）"
    table.cell(2, 1).text = str(c2_count)
    table.cell(3, 0).text = "W（武器卡池）"
    table.cell(3, 1).text = str(w_count)

    heading = doc.add_heading('3. 基础统计', level=1)
    for run in heading.runs:
        set_run_font(run, '微软雅黑', 14, bold=True)
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Light Grid Accent 1'
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, '微软雅黑', 11)
    data = [
        ("总抽卡次数", state.total_draws),
        ("消耗原石", state.total_draws * 160),
        ("五星总数", total_5star),
        ("四星总数", sum(1 for r in state.records_all if r["星级"] == 4)),
        ("总获得星辉", state.starglitter),
        ("总获得星尘", sum(r["获得星尘"] for r in state.records_all)),
        ("可兑换纠缠之缘", state.starglitter // 5)
    ]
    for i, (label, val) in enumerate(data):
        table.cell(i, 0).text = str(label)
        table.cell(i, 1).text = str(val)

    heading = doc.add_heading('4. 关键指标', level=1)
    for run in heading.runs:
        set_run_font(run, '微软雅黑', 14, bold=True)
    kpi_data = [
        ("角色池平均UP抽数", f"{metrics['avg_up_char']:.1f} 抽"),
        ("武器池平均UP抽数", f"{metrics['avg_up_weapon']:.1f} 抽"),
        ("丝柯克平均真实抽数", f"{metrics['avg_skk']:.1f} 抽"),
        ("爱可菲平均真实抽数", f"{metrics['avg_akf']:.1f} 抽"),
        ("苍耀平均真实抽数", f"{metrics['avg_w1']:.1f} 抽"),
        ("香韵奏者平均真实抽数", f"{metrics['avg_w2']:.1f} 抽"),
        ("最多连续UP次数", str(metrics['max_consecutive_up'])),
        ("捕获明光次数", str(metrics['capture_total'])),
        ("提前触发捕获明光", str(metrics['early_capture'])),
        ("小保底不歪概率", f"{metrics['no_loss_rate']:.1%}"),
        ("提前金次数", str(metrics['early_gold'])),
        ("连金次数", str(metrics['consecutive_gold'])),
        ("十连多金次数", str(metrics['multi_gold_10'])),
        ("限定角色真实抽数", f"{metrics['real_up_char']:.1f} 抽"),
        ("限定武器真实抽数", f"{metrics['real_up_weapon']:.1f} 抽"),
    ]
    table = doc.add_table(rows=len(kpi_data), cols=2)
    table.style = 'Light Grid Accent 1'
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, '微软雅黑', 11)
    for i, (label, val) in enumerate(kpi_data):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = val

    heading = doc.add_heading('5. 目标与预测分析', level=1)
    for run in heading.runs:
        set_run_font(run, '微软雅黑', 14, bold=True)
    p = doc.add_paragraph(f"目标预设情况：{'是' if any(state.targets.values()) else '否'}")
    set_run_font(p.runs[0], '微软雅黑', 11)
    if any(state.targets.values()):
        target_desc = []
        if state.targets['C1']:
            target_desc.append(f"丝柯克 {state.targets['C1']} 个")
        if state.targets['C2']:
            target_desc.append(f"爱可菲 {state.targets['C2']} 个")
        if state.targets['W1']:
            target_desc.append(f"苍耀 {state.targets['W1']} 把")
        if state.targets['W2']:
            target_desc.append(f"香韵奏者 {state.targets['W2']} 把")
        p = doc.add_paragraph("目标详情：" + "，".join(target_desc))
        set_run_font(p.runs[0], '微软雅黑', 11)
        if expected_total > 0:
            p = doc.add_paragraph(f"预计期望抽数：约 {expected_total:.0f} 抽")
            set_run_font(p.runs[0], '微软雅黑', 11)
            p = doc.add_paragraph(f"实际消耗抽数：{state.total_draws} 抽")
            set_run_font(p.runs[0], '微软雅黑', 11)
            bias = state.total_draws - expected_total
            bias_percent = (bias / expected_total) * 100 if expected_total else 0
            p = doc.add_paragraph(f"偏差：{bias:+.0f} 抽 ({bias_percent:+.1f}%)")
            set_run_font(p.runs[0], '微软雅黑', 11)
        p = doc.add_paragraph(f"目标达成情况：{'✅ 完全达成' if metrics['target_achieved'] else '❌ 未完全达成'}")
        set_run_font(p.runs[0], '微软雅黑', 11)
        if state.strategy_log:
            p = doc.add_paragraph("策略执行路径：")
            set_run_font(p.runs[0], '微软雅黑', 11, bold=True)
            for log in state.strategy_log:
                log_type, ts, arg1, arg2 = log
                if log_type == "continue":
                    desc = t("weapon.log_continue")
                elif log_type == "stop":
                    desc = t("weapon.log_stop")
                elif log_type == "change":
                    desc = t("weapon.log_change", old=arg1, new=arg2)
                elif log_type == "strategy":
                    desc = t("weapon.log_strategy", weapon=arg1)
                elif log_type == "keep":
                    desc = "保持定轨继续"
                else:
                    desc = str(log)
                p = doc.add_paragraph(f"  [{ts}] {desc}", style='List Bullet')
                for run in p.runs:
                    set_run_font(run, '微软雅黑', 11)

    heading = doc.add_heading('6. 角色池详情', level=1)
    for run in heading.runs:
        set_run_font(run, '微软雅黑', 14, bold=True)
    skk_count = state.char_count.get("丝柯克", 0)
    akf_count = state.char_count.get("爱可菲", 0)
    p = doc.add_paragraph(f"丝柯克：{skk_count} 个（{max(0, skk_count-1)} 命）" + (f"，溢出 {skk_count-7} 命" if skk_count > 7 else ""))
    set_run_font(p.runs[0], '微软雅黑', 11)
    p = doc.add_paragraph(f"爱可菲：{akf_count} 个（{max(0, akf_count-1)} 命）" + (f"，溢出 {akf_count-7} 命" if akf_count > 7 else ""))
    set_run_font(p.runs[0], '微软雅黑', 11)

    std_chars = Counter()
    for r in state.records_char:
        if r["星级"] == 5 and r["获得物品"] not in ["丝柯克", "爱可菲"]:
            std_chars[r["获得物品"]] += 1
    if std_chars:
        p = doc.add_paragraph("常驻五星角色：")
        set_run_font(p.runs[0], '微软雅黑', 11, bold=True)
        for name, cnt in std_chars.items():
            p = doc.add_paragraph(f"  {name}：{cnt} 个", style='List Bullet')
            for run in p.runs:
                set_run_font(run, '微软雅黑', 11)

    heading = doc.add_heading('7. 武器池详情', level=1)
    for run in heading.runs:
        set_run_font(run, '微软雅黑', 14, bold=True)
    w1_count = sum(1 for r in state.records_weapon if r["获得物品"] == t("banner.w1_name"))
    w2_count = sum(1 for r in state.records_weapon if r["获得物品"] == t("banner.w2_name"))
    p = doc.add_paragraph(f"苍耀：{w1_count} 把（精炼 {w1_count} 阶）")
    set_run_font(p.runs[0], '微软雅黑', 11)
    p = doc.add_paragraph(f"香韵奏者：{w2_count} 把（精炼 {w2_count} 阶）")
    set_run_font(p.runs[0], '微软雅黑', 11)

    std_weapons = Counter()
    for r in state.records_weapon:
        if r["星级"] == 5 and r["获得物品"] not in [t("banner.w1_name"), t("banner.w2_name")]:
            std_weapons[r["获得物品"]] += 1
    if std_weapons:
        p = doc.add_paragraph("常驻五星武器：")
        set_run_font(p.runs[0], '微软雅黑', 11, bold=True)
        for name, cnt in std_weapons.items():
            p = doc.add_paragraph(f"  {name}：{cnt} 把", style='List Bullet')
            for run in p.runs:
                set_run_font(run, '微软雅黑', 11)

    if state.records_weapon:
        fate_used = sum(1 for r in state.records_weapon if r.get("命定值") == 1 and r["星级"] == 5)
        p = doc.add_paragraph(f"定轨保底触发次数：{fate_used}")
        set_run_font(p.runs[0], '微软雅黑', 11)

    # 限定物品获取详情（基于真实抽数）
    heading = doc.add_heading('8. 限定物品获取详情', level=1)
    for run in heading.runs:
        set_run_font(run, '微软雅黑', 14, bold=True)
    for item, spins in metrics['real_spins_limited'].items():
        if spins:
            spins_str = "、".join(str(s) for s in spins)
            p = doc.add_paragraph(f"{item}：{spins_str} 抽", style='List Bullet')
        else:
            p = doc.add_paragraph(f"{item}：尚未获得", style='List Bullet')
        for run in p.runs:
            set_run_font(run, '微软雅黑', 11)

    heading = doc.add_heading('9. 四星物品详情', level=1)
    for run in heading.runs:
        set_run_font(run, '微软雅黑', 14, bold=True)
    four_star_chars = Counter()
    four_star_weapons = Counter()
    for r in state.records_all:
        if r["星级"] == 4:
            if r["类别"] == "角色":
                four_star_chars[r["获得物品"]] += 1
            else:
                four_star_weapons[r["获得物品"]] += 1

    if four_star_chars:
        p = doc.add_paragraph("四星角色获取数量（按降序）：")
        set_run_font(p.runs[0], '微软雅黑', 11, bold=True)
        for name, cnt in four_star_chars.most_common(10):
            p = doc.add_paragraph(f"  {name}：{cnt} 个", style='List Bullet')
            for run in p.runs:
                set_run_font(run, '微软雅黑', 11)
    else:
        p = doc.add_paragraph("未获得任何四星角色。")
        set_run_font(p.runs[0], '微软雅黑', 11)

    if four_star_weapons:
        p = doc.add_paragraph("\n四星武器获取数量（按降序）：")
        set_run_font(p.runs[0], '微软雅黑', 11, bold=True)
        for name, cnt in four_star_weapons.most_common(10):
            p = doc.add_paragraph(f"  {name}：{cnt} 把", style='List Bullet')
            for run in p.runs:
                set_run_font(run, '微软雅黑', 11)
    else:
        p = doc.add_paragraph("未获得任何四星武器。")
        set_run_font(p.runs[0], '微软雅黑', 11)

    if chart_files:
        heading = doc.add_heading('10. 可视化图表', level=1)
        for run in heading.runs:
            set_run_font(run, '微软雅黑', 14, bold=True)
        for cf in chart_files:
            if cf and os.path.exists(cf):
                doc.add_picture(cf, width=Inches(5))
                doc.add_paragraph("")

    word_file = excel_file.replace(".xlsx", "_报告.docx")
    doc.save(word_file)
    print(t("export.report_saved", filename=word_file))
    return word_file

# ==================== 14. 主交互循环 ====================
def main():
    print("=" * 60)
    print(t("welcome.title"))
    print("=" * 60)
    print(t("welcome.speed_mode"))
    print(t("welcome.speed_1"))
    print(t("welcome.speed_2"))
    mode_choice = input(t("welcome.input_choice")).strip()
    no_delay = (mode_choice == '2')

    targets, dual_weapon = setup_targets()

    session_time = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    folder_name = f"原神抽卡模拟_{session_time}"
    output_dir = os.path.join("output", folder_name)
    os.makedirs(output_dir, exist_ok=True)
    print(t("output_dir", dir=output_dir))

    print("=" * 60)
    print(t("banner.info_title"))
    print(t("banner.c1_info", ups=", ".join(FOUR_STAR_CHARS_UP)))
    print(t("banner.c2_info", ups=", ".join(FOUR_STAR_CHARS_UP)))
    print(t("banner.w_info", ups=", ".join(FOUR_STAR_WEAPONS_UP)))
    print(t("banner.note_weapon"))
    print(t("banner.note_char"))
    print(t("banner.note_four"))
    print(t("banner.note_weapon_pause"))
    print("-" * 60)
    print(t("command.instructions"))
    print(t("command.pull_count"))
    print(t("command.format"))
    print(t("command.c1_example"))
    print(t("command.c2_example"))
    print(t("command.w1_example"))
    print(t("command.w2_example"))
    print(t("command.y_repeat"))
    print(t("command.s_end"))
    print("=" * 60)

    state = GachaState()
    state.targets = targets
    state.dual_weapon_mode = dual_weapon

    last_banner_type = None
    last_banner_code = None
    last_weapon_choice = None
    last_count = 0

    disabled_pools = set()

    while True:
        cmd = input(t("command.enter")).strip()
        if cmd.upper() == 'S':
            break
        elif cmd.upper() == 'Y':
            if last_banner_type is None:
                print(t("command.no_previous"))
                continue
            banner_type = last_banner_type
            banner_code = last_banner_code
            weapon_choice = last_weapon_choice
            count = last_count
            print(t("command.repeat", count=count))
        else:
            parts = cmd.split()
            if len(parts) < 2:
                print(t("command.invalid_format"))
                continue
            banner_code_input = parts[0].upper()
            try:
                count = int(parts[1])
                if count <= 0:
                    print(t("command.invalid_count"))
                    continue
            except ValueError:
                print(t("command.invalid_number"))
                continue

            if banner_code_input in ['C1', 'C2']:
                if banner_code_input in disabled_pools:
                    print(t("command.disabled_pool", banner=banner_code_input))
                    continue
                banner_type = 'char'
                banner_code = banner_code_input
                weapon_choice = None
            elif banner_code_input == 'W':
                if len(parts) == 2:
                    weapon_choice = None
                else:
                    weapon_code = parts[2].upper()
                    if weapon_code not in ['W1', 'W2']:
                        print(t("command.invalid_banner"))
                        continue
                    weapon_choice = weapon_code
                banner_type = 'weapon'
                banner_code = 'W'
            else:
                print(t("command.invalid_banner"))
                continue

            last_banner_type = banner_type
            last_banner_code = banner_code
            last_weapon_choice = weapon_choice
            last_count = count

        pool_disabled = perform_draws(state, count, banner_type, banner_code, weapon_choice, no_delay, disabled_pools)
        if pool_disabled:
            pass
        print_limited_summary(state)
        print_target_status(state)
        print(t("command.completed"))
        print(t("command.hint"))

    if not state.records_all:
        print(t("no_pulls"))
        return

    print(t("export.title"))
    print(t("export.option1"))
    print(t("export.option2"))
    print(t("export.option3"))
    print(t("export.option4"))
    choice = input(t("export.enter")).strip()

    expected_total = 0
    if state.targets['C1'] > 0:
        expected_total += EXPECTATION['c6_char'] if state.targets['C1'] == 7 else EXPECTATION['single_up_char'] * state.targets['C1']
    if state.targets['C2'] > 0:
        expected_total += EXPECTATION['c6_char'] if state.targets['C2'] == 7 else EXPECTATION['single_up_char'] * state.targets['C2']
    if state.targets['W1'] > 0:
        expected_total += EXPECTATION['single_up_weapon'] * state.targets['W1']
    if state.targets['W2'] > 0:
        expected_total += EXPECTATION['single_up_weapon'] * state.targets['W2']

    excel_file = save_to_excel(state, output_dir, expected_total)

    if choice in ['2', '4']:
        chart_files = generate_plots(state, excel_file)
    else:
        chart_files = None

    if choice in ['3', '4']:
        if choice == '3':
            chart_files = generate_plots(state, excel_file)
        word_file = generate_word_report(state, excel_file, chart_files, expected_total)

    print(t("export.completed", output_dir=output_dir))
    print(t("goodbye"))

if __name__ == "__main__":
    main()